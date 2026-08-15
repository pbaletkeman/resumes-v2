"""
renderer.py
Template-based multi-format resume renderer.

Renders ``RewriteOutput`` and ``CoverLetterOutput`` models against
Jinja2 templates to produce plaintext, Markdown, DOCX (python-docx), and
PDF (ReportLab Platypus) output.  DOCX/PDF rendering is fully supported:
``render_docx`` / ``render_pdf`` and ``render_cover_letter_docx`` /
``render_cover_letter_pdf`` build the documents directly from the same
context the text templates use (see :meth:`ResumeRenderer.render_all`).

Rendering path: this is *the* template-based path.  ``ResumeRenderer``
loads Jinja2 sources from ``client.templates`` (``TEMPLATES`` +
``COVER_LETTER``), builds a context dict from the Pydantic models
(:meth:`ResumeRenderer._build_context` /
:meth:`ResumeRenderer._build_cover_letter_context`), and renders every
format (text, DOCX, PDF) from that one context.  The alternative,
simpler path lives in ``client/formatter.py``: plain string-building
helpers (``format_resume_markdown`` / ``format_resume_plain`` /
``format_cover_letter``) with no templates and no DOCX/PDF support.
Prefer ``ResumeRenderer`` for multi-format output; use the formatter
helpers for a single plain/markdown string.
"""

import os
import re
import tempfile
from datetime import date, datetime
from pathlib import Path
from typing import Any, cast
from xml.sax.saxutils import escape

from docx import Document as create_document
from docx.document import Document as DocumentType
from docx.shared import Inches, Pt
from docx.styles.style import ParagraphStyle
from jinja2 import BaseLoader, Environment, StrictUndefined

try:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle as PdfParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Flowable,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )
except ImportError as exc:  # pragma: no cover
    raise ImportError(
        "ReportLab is required for PDF rendering. Install it with "
        "`uv sync` (adds `reportlab>=4.0` to the project dependencies)."
    ) from exc

from client.models import CoverLetterOutput, RewriteOutput
from client.templates import TEMPLATES
from client.templates.cover_letter import COVER_LETTER


class ResumeRenderer:
    """Renders resume data using Jinja2 templates.

    Templates are loaded from the ``client.templates`` package by default.
    Pass *template_dir* to load from a custom directory instead.

    Args:
        template_dir: Optional directory to load ``.j2`` template files from.
            When ``None`` (the default), uses the built-in template dicts
            from ``client.templates``.
    """

    def __init__(self, template_dir: Path | None = None) -> None:
        self._template_dir = template_dir
        self._env = Environment(
            loader=BaseLoader(),
            undefined=StrictUndefined,
            keep_trailing_newline=True,
        )
        # Built-in templates from the package
        self._templates: dict[str, dict[str, str]] = dict(TEMPLATES)

    def render_plaintext(
        self,
        resume: RewriteOutput,
        *,
        name: str = "",
        title: str = "",
        template: str = "modern",
    ) -> str:
        """Render *resume* as clean plaintext using the named template.

        Args:
            resume: Structured resume data from the Resume Rewrite Agent.
            name: Candidate name for the header.
            title: Candidate title for the header.
            template: Template key (``"modern"``, ``"classic"``, or
                ``"minimal"``).

        Returns:
            A rendered plaintext string.

        Raises:
            KeyError: If *template* is not found.
            jinja2.UndefinedError: If the template references a variable
                that is not provided.
        """
        tpl_dict = self._templates[template]
        tpl_source = tpl_dict["plaintext"]

        context = self._build_context(resume, name=name, title=title)
        return self._render(tpl_source, context)

    def render_markdown(
        self,
        resume: RewriteOutput,
        *,
        name: str = "",
        title: str = "",
        template: str = "modern",
    ) -> str:
        """Render *resume* as Markdown using the named template.

        Args:
            resume: Structured resume data from the Resume Rewrite Agent.
            name: Candidate name for the header.
            title: Candidate title for the header.
            template: Template key (``"modern"``, ``"classic"``, or
                ``"minimal"``).

        Returns:
            A rendered Markdown string.

        Raises:
            KeyError: If *template* is not found.
            jinja2.UndefinedError: If the template references a variable
                that is not provided.
        """
        tpl_dict = self._templates[template]
        tpl_source = tpl_dict["markdown"]

        context = self._build_context(resume, name=name, title=title)
        return self._render(tpl_source, context)

    def render_cover_letter_plaintext(
        self,
        cover_letter: CoverLetterOutput,
        *,
        name: str = "",
        company: str = "",
        phone: str = "",
        email: str = "",
        linkedin: str = "",
        github: str = "",
    ) -> str:
        """Render *cover_letter* as clean plaintext.

        Args:
            cover_letter: Structured letter data from the Cover Letter Agent.
            name: Candidate name for the signature and header.
            company: Target company name (reserved for context).
            phone: Candidate phone number for the header.
            email: Candidate email address for the header.
            linkedin: Candidate LinkedIn profile URL for the header.
            github: Candidate GitHub profile URL for the header.

        Returns:
            A rendered plaintext letter string.

        Raises:
            jinja2.UndefinedError: If the template references a variable
                that is not provided.
        """
        tpl_source = COVER_LETTER["plaintext"]

        context = self._build_cover_letter_context(
            cover_letter,
            name=name,
            company=company,
            phone=phone,
            email=email,
            linkedin=linkedin,
            github=github,
        )
        return self._render(tpl_source, context)

    def render_cover_letter_markdown(
        self,
        cover_letter: CoverLetterOutput,
        *,
        name: str = "",
        company: str = "",
        phone: str = "",
        email: str = "",
        linkedin: str = "",
        github: str = "",
    ) -> str:
        """Render *cover_letter* as Markdown.

        Args:
            cover_letter: Structured data from the Cover Letter Agent.
            name: Candidate name for the signature and header.
            company: Target company name (reserved for context).
            phone: Candidate phone number for the header.
            email: Candidate email address for the header.
            linkedin: Candidate LinkedIn profile URL for the header.
            github: Candidate GitHub profile URL for the header.

        Returns:
            A rendered Markdown letter string.

        Raises:
            jinja2.UndefinedError: If the template references a variable
                that is not provided.
        """
        tpl_source = COVER_LETTER["markdown"]

        context = self._build_cover_letter_context(
            cover_letter,
            name=name,
            company=company,
            phone=phone,
            email=email,
            linkedin=linkedin,
            github=github,
        )
        return self._render(tpl_source, context)

    def render_cover_letter_docx(
        self,
        cover_letter: CoverLetterOutput,
        *,
        name: str = "",
        company: str = "",
        phone: str = "",
        email: str = "",
        linkedin: str = "",
        github: str = "",
        output_path: str | Path | None = None,
    ) -> Path:
        """Render *cover_letter* as a professionally styled ``.docx`` document.

        Letter-size pages with 1-inch margins and Calibri 11pt body text,
        mirroring the plaintext/Markdown letter layout: name header, contact
        line, date, salutation, body paragraphs, and signature.  When
        *output_path* is ``None`` a temporary path is used.

        Args:
            cover_letter: Structured letter data from the Cover Letter Agent.
            name: Candidate name for the header and signature.
            company: Target company name (reserved for context).
            phone: Candidate phone number for the header.
            email: Candidate email address for the header.
            linkedin: Candidate LinkedIn profile URL for the header.
            github: Candidate GitHub profile URL for the header.
            output_path: Destination file. Defaults to a temp file.

        Returns:
            The ``Path`` the document was written to.

        Raises:
            OSError: If the output directory cannot be created or the
                document cannot be saved (disk/permission errors).
        """
        doc: DocumentType = create_document()

        for section in doc.sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        normal = cast(ParagraphStyle, doc.styles["Normal"])
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        context = self._build_cover_letter_context(
            cover_letter,
            name=name,
            company=company,
            phone=phone,
            email=email,
            linkedin=linkedin,
            github=github,
        )
        self._populate_cover_letter_docx(doc, context)

        path = Path(output_path) if output_path is not None else _temp_docx_path()
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return path

    def render_cover_letter_pdf(
        self,
        cover_letter: CoverLetterOutput,
        *,
        name: str = "",
        company: str = "",
        phone: str = "",
        email: str = "",
        linkedin: str = "",
        github: str = "",
        output_path: str | Path | None = None,
    ) -> Path:
        """Render *cover_letter* as a professionally styled PDF document.

        ReportLab builds the PDF directly with Platypus (no HTML/Markdown
        intermediate), mirroring the DOCX letter layout with the shared
        :meth:`_pdf_styles`.  When *output_path* is ``None`` a temporary
        path is used.

        Args:
            cover_letter: Structured letter data from the Cover Letter Agent.
            name: Candidate name for the header and signature.
            company: Target company name (reserved for context).
            phone: Candidate phone number for the header.
            email: Candidate email address for the header.
            linkedin: Candidate LinkedIn profile URL for the header.
            github: Candidate GitHub profile URL for the header.
            output_path: Destination file. Defaults to a temp file.

        Returns:
            The ``Path`` the document was written to.

        Raises:
            OSError: If the output directory cannot be created or the
                PDF cannot be built (disk/permission errors).
        """
        path = Path(output_path) if output_path is not None else _temp_pdf_path()
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(
            str(path),
            pagesize=letter,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )
        context = self._build_cover_letter_context(
            cover_letter,
            name=name,
            company=company,
            phone=phone,
            email=email,
            linkedin=linkedin,
            github=github,
        )
        flowables = self._populate_cover_letter_pdf_flowables(
            context, self._pdf_styles()
        )
        doc.build(flowables)
        return path

    def render_docx(
        self,
        resume: RewriteOutput,
        *,
        name: str = "",
        title: str = "",
        template: str = "modern",
        output_path: str | Path | None = None,
    ) -> Path:
        """Render *resume* as a professionally styled ``.docx`` document.

        The document uses letter-size pages with 1-inch margins, Calibri
        11pt body text, and the name rendered large and bold.  When
        *output_path* is ``None`` a temporary path is used.

        Args:
            resume: Structured resume data from the Resume Rewrite Agent.
            name: Candidate name for the header.
            title: Candidate title for the header.
            template: Accepted for API consistency (styling is fixed).
            output_path: Destination file. Defaults to a temp file.

        Returns:
            The ``Path`` the document was written to.

        Raises:
            OSError: If the output directory cannot be created or the
                document cannot be saved (disk/permission errors).
        """
        doc: DocumentType = create_document()

        for section in doc.sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        normal = cast(ParagraphStyle, doc.styles["Normal"])
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        context = self._build_context(resume, name=name, title=title)
        self._populate_docx_paragraphs(doc, context)

        path = Path(output_path) if output_path is not None else _temp_docx_path()
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return path

    def render_pdf(
        self,
        resume: RewriteOutput,
        *,
        name: str = "",
        title: str = "",
        template: str = "modern",
        output_path: str | Path | None = None,
    ) -> Path:
        """Render *resume* as a professionally styled PDF document.

        ReportLab builds the PDF directly with Platypus (no HTML/Markdown
        intermediate).  Letter-size pages with 1-inch margins and the
        shared :meth:`_pdf_styles` styling.  When *output_path* is ``None``
        a temporary path is used.

        Args:
            resume: Structured resume data from the Resume Rewrite Agent.
            name: Candidate name for the header.
            title: Candidate title for the header.
            template: Accepted for API consistency (styling is fixed).
            output_path: Destination file. Defaults to a temp file.

        Returns:
            The ``Path`` the document was written to.

        Raises:
            OSError: If the output directory cannot be created or the
                PDF cannot be built (disk/permission errors).
        """
        path = Path(output_path) if output_path is not None else _temp_pdf_path()
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(
            str(path),
            pagesize=letter,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )
        context = self._build_context(resume, name=name, title=title)
        flowables = self._populate_pdf_flowables(context, self._pdf_styles())
        doc.build(flowables)
        return path

    def render_all(
        self,
        resume: RewriteOutput,
        cover_letter: CoverLetterOutput | None,
        *,
        candidate_name: str,
        company_name: str,
        output_dir: str | Path,
        resume_template: str = "modern",
        resume_templates: str | list[str] | None = None,
        phone: str = "",
        email: str = "",
        linkedin: str = "",
        github: str = "",
    ) -> dict[str, Path]:
        """Render *resume* and, when available, *cover_letter* into formats.

        Produces the four resume formats (plaintext, Markdown, DOCX, PDF).
        When *cover_letter* is provided and non-empty, the four cover letter
        formats (plaintext, Markdown, DOCX, PDF) are produced as well;
        otherwise they are skipped.  Each file is written to *output_dir*
        under a timestamped, slugified filename built by
        :meth:`build_output_path`.

        By default a single resume template (``resume_template``, "modern")
        is rendered and its files use the ``resume_plaintext`` ...
        ``resume_pdf`` keys.  Pass *resume_templates* (a template key or a
        list of keys) to render several resume layouts in one call instead;
        their files are namespaced as ``resume_{template}_plaintext`` ...
        ``resume_{template}_pdf`` and the filename embeds the template
        (``resume-{template}.{ext}``) so the layouts do not overwrite each
        other.

        Args:
            resume: Structured resume data from the Resume Rewrite Agent.
            cover_letter: Structured letter data from the Cover Letter Agent,
                or ``None`` to skip the letter formats.
            candidate_name: Candidate name for headers and filenames.
            company_name: Target company name for filenames.
            output_dir: Directory the rendered files are written to.
            resume_template: Template key for the single-template resume
                formats, used when *resume_templates* is ``None``.
            resume_templates: Optional template key or list of keys
                (``"modern"``/``"classic"``/``"minimal"``) to render multiple
                resume layouts in one call.  When provided, each layout's
                files are namespaced with the template embedded in the key
                and filename.
            phone: Candidate phone number for cover letter headers.
            email: Candidate email address for cover letter headers.
            linkedin: Candidate LinkedIn profile URL for cover letter headers.
            github: Candidate GitHub profile URL for cover letter headers.

        Returns:
            Mapping of format name to the written ``Path``.

        Raises:
            KeyError: If *resume_template*/*resume_templates* references a
                template key that is not built in (``"modern"``,
                ``"classic"``, ``"minimal"``).
            jinja2.UndefinedError: If a template references a variable
                that is not provided in the built context.
            OSError: If *output_dir* cannot be created or a rendered
                file cannot be written (disk/permission errors).
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        multi = resume_templates is not None
        if resume_templates is None:
            templates: list[str] = [resume_template]
        elif isinstance(resume_templates, str):
            templates = [resume_templates]
        else:
            templates = list(resume_templates)

        paths: dict[str, Path] = {}
        for template in templates:
            paths.update(
                self._render_resume_template(
                    resume,
                    candidate_name=candidate_name,
                    company_name=company_name,
                    output_dir=output_dir,
                    template=template,
                    key_prefix=f"resume_{template}" if multi else "resume",
                    document_type=f"resume-{template}" if multi else "resume",
                )
            )

        if cover_letter is not None and cover_letter.cover_letter.strip():
            letter_plain = self.render_cover_letter_plaintext(
                cover_letter,
                name=candidate_name,
                company=company_name,
                phone=phone,
                email=email,
                linkedin=linkedin,
                github=github,
            )
            letter_md = self.render_cover_letter_markdown(
                cover_letter,
                name=candidate_name,
                company=company_name,
                phone=phone,
                email=email,
                linkedin=linkedin,
                github=github,
            )
            letter_docx = self.render_cover_letter_docx(
                cover_letter,
                name=candidate_name,
                company=company_name,
                phone=phone,
                email=email,
                linkedin=linkedin,
                github=github,
                output_path=self.build_output_path(
                    "cover_letter",
                    candidate_name=candidate_name,
                    company_name=company_name,
                    output_dir=output_dir,
                    ext=".docx",
                ),
            )
            letter_pdf = self.render_cover_letter_pdf(
                cover_letter,
                name=candidate_name,
                company=company_name,
                phone=phone,
                email=email,
                linkedin=linkedin,
                github=github,
                output_path=self.build_output_path(
                    "cover_letter",
                    candidate_name=candidate_name,
                    company_name=company_name,
                    output_dir=output_dir,
                    ext=".pdf",
                ),
            )
            paths["cover_letter_plaintext"] = self._write_text(
                letter_plain,
                "cover_letter",
                candidate_name,
                company_name,
                output_dir,
                ".txt",
            )
            paths["cover_letter_markdown"] = self._write_text(
                letter_md,
                "cover_letter",
                candidate_name,
                company_name,
                output_dir,
                ".md",
            )
            paths["cover_letter_docx"] = letter_docx
            paths["cover_letter_pdf"] = letter_pdf

        return paths

    def _render_resume_template(
        self,
        resume: RewriteOutput,
        *,
        candidate_name: str,
        company_name: str,
        output_dir: Path,
        template: str,
        key_prefix: str,
        document_type: str,
    ) -> dict[str, Path]:
        """Render *resume* in *template* into its four output formats.

        The returned mapping uses the ``{key_prefix}_plaintext`` /
        ``{key_prefix}_markdown`` / ``{key_prefix}_docx`` /
        ``{key_prefix}_pdf`` keys.  Text files are written under
        *document_type* so ``build_output_path`` produces distinct filenames
        per template (e.g. ``resume-classic.md``).

        Args:
            resume: Structured resume data from the Resume Rewrite Agent.
            candidate_name: Candidate name for headers and filenames.
            company_name: Target company name for filenames.
            output_dir: Directory the rendered files are written to.
            template: Template key (``"modern"``, ``"classic"``, or
                ``"minimal"``).
            key_prefix: Prefix for the four output keys.
            document_type: Document type used for the output filenames.

        Returns:
            Mapping of the four format keys to written ``Path`` values.

        Raises:
            KeyError: If *template* is not a built-in template key.
        """
        plain = self.render_plaintext(resume, name=candidate_name, template=template)
        md = self.render_markdown(resume, name=candidate_name, template=template)
        docx = self.render_docx(
            resume,
            name=candidate_name,
            template=template,
            output_path=self.build_output_path(
                document_type,
                candidate_name=candidate_name,
                company_name=company_name,
                output_dir=output_dir,
                ext=".docx",
            ),
        )
        pdf = self.render_pdf(
            resume,
            name=candidate_name,
            template=template,
            output_path=self.build_output_path(
                document_type,
                candidate_name=candidate_name,
                company_name=company_name,
                output_dir=output_dir,
                ext=".pdf",
            ),
        )
        return {
            f"{key_prefix}_plaintext": self._write_text(
                plain, document_type, candidate_name, company_name, output_dir, ".txt"
            ),
            f"{key_prefix}_markdown": self._write_text(
                md, document_type, candidate_name, company_name, output_dir, ".md"
            ),
            f"{key_prefix}_docx": docx,
            f"{key_prefix}_pdf": pdf,
        }

    @staticmethod
    def _write_text(
        content: str,
        document_type: str,
        candidate_name: str,
        company_name: str,
        output_dir: Path,
        ext: str,
    ) -> Path:
        """Write *content* to a timestamped output file and return its path.

        Args:
            content: Text to write (UTF-8).
            document_type: Type of document, e.g. ``"resume"`` or
                ``"cover_letter"`` (used for the filename and extension).
            candidate_name: Candidate name for the filename segment.
            company_name: Company name for the filename segment.
            output_dir: Directory the file is written to.
            ext: File extension including the leading dot.

        Returns:
            The ``Path`` of the written file.

        Raises:
            OSError: If *output_dir* cannot be created or *content*
                cannot be written (disk/permission errors).
        """
        path = ResumeRenderer.build_output_path(
            document_type,
            candidate_name=candidate_name,
            company_name=company_name,
            output_dir=output_dir,
            ext=ext,
        )
        path.write_text(content, encoding="utf-8")
        return path

    @staticmethod
    def build_output_path(
        document_type: str,
        *,
        candidate_name: str,
        company_name: str,
        output_dir: str | Path,
        ext: str | None = None,
    ) -> Path:
        """Build a timestamped output file path for *document_type*.

        The returned path has the form::

            {output_dir}/{YYYYMMDD}_{candidate}_{company6}_{document_type}.{ext}

        where ``company6`` is the first six characters of the company name.
        The date carries no time so a day's runs collide (files are
        overwritten) rather than accumulating minute-by-minute variants.
        Names are slugified (see :meth:`_slugify`) so the result is safe
        to write on any filesystem.  Pure path logic -- no file I/O.

        Args:
            document_type: Type of document, e.g. ``"resume"`` or
                ``"cover_letter"``.
            candidate_name: Candidate name for the path segment.
            company_name: Company name for the path segment.
            output_dir: Directory the file will live in.
            ext: File extension including the leading dot.  When ``None``,
                a default is chosen per *document_type*.

        Returns:
            The resolved ``Path`` under *output_dir*.
        """
        resolved_ext = ext or _default_extension(document_type)
        if not resolved_ext.startswith("."):
            resolved_ext = f".{resolved_ext}"

        timestamp = datetime.now().strftime("%Y%m%d")
        slug_candidate = ResumeRenderer._slugify(candidate_name)
        slug_company = ResumeRenderer._slugify(company_name[:6])
        slug_type = ResumeRenderer._slugify(document_type) or "output"

        filename = (
            f"{timestamp}_{slug_candidate}_{slug_company}_{slug_type}{resolved_ext}"
        )
        return Path(output_dir) / filename

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _render(self, template_source: str, context: dict[str, object]) -> str:
        """Render *context* against *template_source* and clean the output.

        Shared by every text-format public method: run the Jinja2 render
        then collapse excess blank lines via :meth:`_clean_output`.  The
        caller is responsible for resolving *template_source* from the
        right template container (``self._templates[template][fmt]`` for
        resume formats, ``COVER_LETTER[fmt]`` for letters) and building
        *context* with :meth:`_build_context` /
        :meth:`_build_cover_letter_context`.

        Args:
            template_source: A Jinja2 template string.
            context: Template variables, keyed by name.

        Returns:
            The cleaned rendered output.

        Raises:
            jinja2.UndefinedError: If the template references a variable
                that is not provided.
        """
        rendered = self._env.from_string(template_source).render(**context)
        return self._clean_output(rendered)

    @staticmethod
    def _slugify(text: str) -> str:
        """Normalize *text* to a filename-safe ASCII token.

        Lowercases the input, replaces runs of non-alphanumeric characters
        with a single ``-``, and strips leading/trailing hyphens.  Returns
        an empty string when nothing usable remains.
        """
        ascii_text = text.encode("ascii", errors="ignore").decode("ascii")
        slug = re.sub(r"[^a-z0-9]+", "-", ascii_text.lower()).strip("-")
        return slug

    @staticmethod
    def _build_context(
        resume: RewriteOutput,
        *,
        name: str = "",
        title: str = "",
    ) -> dict[str, object]:
        """Convert a ``RewriteOutput`` into a Jinja2 template context dict.

        Args:
            resume: Structured resume data from the Resume Rewrite Agent.
            name: Candidate name for the header.
            title: Candidate title for the header.

        Returns:
            A context dict with the keys the resume templates expect:
            ``name``, ``title``, ``summary``, ``skills``, ``experience``
            (list of per-job dicts), ``projects``, ``certifications``,
            and ``education``.
        """
        return {
            "name": name,
            "title": title,
            "summary": resume.summary,
            "skills": resume.skills,
            "experience": [
                {
                    "title": job.title,
                    "company": job.company,
                    "dates": job.dates,
                    "responsibilities": job.responsibilities,
                    "achievements": job.achievements,
                    "metrics": job.metrics,
                }
                for job in resume.experience
            ],
            "projects": resume.projects,
            "certifications": resume.certifications,
            "education": resume.education,
        }

    @staticmethod
    def _build_cover_letter_context(
        cover_letter: CoverLetterOutput,
        *,
        name: str = "",
        company: str = "",
        phone: str = "",
        email: str = "",
        linkedin: str = "",
        github: str = "",
    ) -> dict[str, object]:
        """Convert a ``CoverLetterOutput`` into a Jinja2 context dict.

        The ``COVER_LETTER`` template expects ``candidate_name``, ``date``,
        ``opening_paragraph``, ``body_paragraph``, and ``closing_paragraph``.
        The letter body is split on blank lines into up to three logical
        paragraphs (opening / middle / closing).  ``contact_line`` holds the
        non-empty contact details joined on `` | `` for the header.

        The body's own closing signature is dropped so the template-supplied
        one is not duplicated; any contact details that rode along with that
        signature (appended by ``_apply_contact_info`` or the fallback letter)
        are merged into ``contact_line`` instead of being lost.
        """
        paragraphs, trailing_contact = _split_letter_body(cover_letter.cover_letter)
        opening = paragraphs[0] if paragraphs else ""
        closing = paragraphs[-1] if len(paragraphs) > 1 else ""
        body = "\n\n".join(paragraphs[1:-1]) if len(paragraphs) > 2 else ""
        header_parts = [p for p in (phone, email, linkedin, github) if p]
        if trailing_contact:
            header_parts += [
                p.strip() for p in trailing_contact.split(" | ") if p.strip()
            ]
        contact_parts: list[str] = []
        for part in header_parts:
            if part not in contact_parts:
                contact_parts.append(part)
        return {
            "candidate_name": name,
            "company": company,
            "phone": phone,
            "email": email,
            "linkedin": linkedin,
            "github": github,
            "contact_line": " | ".join(contact_parts),
            "date": date.today().strftime("%B %d, %Y"),
            "opening_paragraph": opening,
            "body_paragraph": body,
            "closing_paragraph": closing,
        }

    @staticmethod
    def _clean_output(text: str) -> str:
        """Collapse excessive blank lines produced by Jinja2 templates.

        Args:
            text: Raw rendered output.

        Returns:
            The rendered text with runs of blank lines reduced to a single
            blank line, trailing whitespace stripped, and the outer
            whitespace removed.
        """
        lines = text.split("\n")
        cleaned: list[str] = []
        prev_blank = False
        for line in lines:
            stripped = line.rstrip()
            is_blank = stripped == ""
            if is_blank and prev_blank:
                continue
            cleaned.append(stripped)
            prev_blank = is_blank
        return "\n".join(cleaned).strip()

    @staticmethod
    def _docx_heading(doc: DocumentType, text: str) -> None:
        """Add a bold section heading paragraph to *doc*.

        Encapsulates the heading run styling so every section header
        shares the same font size and weight.
        """
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)

    @staticmethod
    def _docx_bullet(doc: DocumentType, text: str) -> None:
        """Add a bulleted paragraph to *doc* using the ``List Bullet`` style.

        Encapsulates the bullet list styling so all list content shares
        the same paragraph style.
        """
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(text)

    @staticmethod
    def _pdf_styles() -> dict[str, PdfParagraphStyle]:
        """Return the shared set of ReportLab ``ParagraphStyle`` objects.

        Uses Helvetica / Helvetica-Bold (Type-1 base-14 fonts), a 14pt bold
        name, bold section headings, and 10.5pt single-spaced body text --
        matching the professional DOCX layout.  Base-14 fonts cannot render
        Unicode/emoji, which is fine given the project's ASCII-only rule.
        """
        return {
            "name": PdfParagraphStyle(
                name="ResumeName",
                fontName="Helvetica-Bold",
                fontSize=14,
                leading=18,
                spaceAfter=2,
                alignment=TA_LEFT,
            ),
            "title": PdfParagraphStyle(
                name="ResumeTitle",
                fontName="Helvetica",
                fontSize=11,
                leading=14,
                spaceAfter=8,
                alignment=TA_LEFT,
            ),
            "heading": PdfParagraphStyle(
                name="SectionHeading",
                fontName="Helvetica-Bold",
                fontSize=12,
                leading=15,
                spaceBefore=10,
                spaceAfter=4,
                alignment=TA_LEFT,
            ),
            "body": PdfParagraphStyle(
                name="Body",
                fontName="Helvetica",
                fontSize=10.5,
                leading=13,
                spaceAfter=4,
                alignment=TA_LEFT,
            ),
            "bullet": PdfParagraphStyle(
                name="Bullet",
                fontName="Helvetica",
                fontSize=10.5,
                leading=13,
                leftIndent=14,
                spaceAfter=2,
                alignment=TA_LEFT,
            ),
        }

    @staticmethod
    def _populate_pdf_flowables(
        context: dict[str, object], styles: dict[str, PdfParagraphStyle]
    ) -> list[Flowable]:
        """Build the Platypus flowables for the resume body in *context*.

        Mirrors :meth:`_populate_docx_paragraphs`: a name/title header,
        summary paragraph, bulleted skills, per-experience blocks with a
        bold title/company header line, and bulleted projects /
        certifications / education.  Text is XML-escaped for ReportLab's
        ``Paragraph`` markup parser (bold via inline ``<b>`` tags).
        """
        flowables: list[Flowable] = []
        name = str(context.get("name", ""))
        title = str(context.get("title", ""))

        if name:
            flowables.append(Paragraph(escape(name), styles["name"]))
        if title:
            flowables.append(Paragraph(escape(title), styles["title"]))

        summary = str(context.get("summary", "")).strip()
        if summary:
            flowables.append(Paragraph("Summary", styles["heading"]))
            flowables.append(Paragraph(escape(summary), styles["body"]))

        skills = cast(list[str], context.get("skills", []))
        if skills:
            flowables.append(Paragraph("Skills", styles["heading"]))
            flowables.append(_bullet_list(skills, styles))

        experience = cast(list[dict[str, object]], context.get("experience", []))
        if experience:
            flowables.append(Paragraph("Experience", styles["heading"]))
            for job in experience:
                job_title = str(job.get("title", ""))
                company = str(job.get("company", ""))
                dates = str(job.get("dates", ""))

                pieces: list[str] = []
                if job_title:
                    pieces.append(f"<b>{escape(job_title)}</b>")
                if company:
                    pieces.append(f"<b>{escape(company)}</b>")
                if dates:
                    pieces.append(escape(dates))
                if pieces:
                    flowables.append(Paragraph(" - ".join(pieces), styles["body"]))

                for key in ("responsibilities", "achievements", "metrics"):
                    items = cast(list[str], job.get(key, []))
                    if items:
                        flowables.append(_bullet_list([text for text in items], styles))
                flowables.append(Spacer(1, 6))

        for section, label in (
            ("projects", "Projects"),
            ("certifications", "Certifications"),
            ("education", "Education"),
        ):
            items = cast(list[str], context.get(section, []))
            if items:
                flowables.append(Paragraph(label, styles["heading"]))
                flowables.append(_bullet_list(items, styles))

        return flowables

    @staticmethod
    def _populate_docx_paragraphs(
        doc: DocumentType, context: dict[str, object]
    ) -> None:
        """Populate *doc* with the resume content in *context*.

        *context* is the dict produced by :meth:`_build_context`.  The
        candidate name is rendered at 14pt bold; section headings are bold;
        experience is rendered as blocks of a bold title and company (with
        plain dates) plus bulleted responsibilities, achievements, and
        metrics.
        """
        name = str(context.get("name", ""))
        title = str(context.get("title", ""))

        if name:
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(name)
            name_run.bold = True
            name_run.font.size = Pt(14)
        if title:
            title_para = doc.add_paragraph()
            title_para.add_run(title)

        summary = str(context.get("summary", "")).strip()
        if summary:
            ResumeRenderer._docx_heading(doc, "Summary")
            doc.add_paragraph(summary)

        skills = cast(list[str], context.get("skills", []))
        if skills:
            ResumeRenderer._docx_heading(doc, "Skills")
            for item in skills:
                ResumeRenderer._docx_bullet(doc, item)

        experience = cast(list[dict[str, object]], context.get("experience", []))
        if experience:
            ResumeRenderer._docx_heading(doc, "Experience")
            for job in experience:
                job_title = str(job.get("title", ""))
                company = str(job.get("company", ""))
                dates = str(job.get("dates", ""))

                if job_title or company or dates:
                    header = doc.add_paragraph()
                    if job_title:
                        title_run = header.add_run(job_title)
                        title_run.bold = True
                    if company:
                        company_run = header.add_run(f" - {company}")
                        company_run.bold = True
                    if dates:
                        header.add_run(f" - {dates}")

                for key in ("responsibilities", "achievements", "metrics"):
                    for item in cast(list[str], job.get(key, [])):
                        ResumeRenderer._docx_bullet(doc, item)

        projects = cast(list[str], context.get("projects", []))
        if projects:
            ResumeRenderer._docx_heading(doc, "Projects")
            for item in projects:
                ResumeRenderer._docx_bullet(doc, item)

        certifications = cast(list[str], context.get("certifications", []))
        if certifications:
            ResumeRenderer._docx_heading(doc, "Certifications")
            for item in certifications:
                ResumeRenderer._docx_bullet(doc, item)

        education = cast(list[str], context.get("education", []))
        if education:
            ResumeRenderer._docx_heading(doc, "Education")
            for item in education:
                ResumeRenderer._docx_bullet(doc, item)

    @staticmethod
    def _cover_letter_paragraphs(
        context: dict[str, object],
    ) -> list[tuple[str, str]]:
        """Return the letter's ordered ``(text, kind)`` paragraphs.

        Mirrors the ``COVER_LETTER`` template layout: header (candidate
        name, contact line, date), salutation, opening/body/closing
        paragraphs, and the closing signature (``Sincerely,`` plus the
        candidate name).  *kind* is one of ``"name"``, ``"meta"``,
        ``"salutation"``, ``"body"``, or ``"signature"`` so the DOCX/PDF
        populate helpers can style each block without re-deriving the
        letter structure.

        Args:
            context: The context dict from :meth:`_build_cover_letter_context`.

        Returns:
            Ordered ``(text, kind)`` pairs ready for DOCX/PDF rendering.
        """
        paragraphs: list[tuple[str, str]] = []
        name = str(context.get("candidate_name", "")).strip()
        contact_line = str(context.get("contact_line", "")).strip()
        date = str(context.get("date", "")).strip()
        hiring_manager = str(context.get("hiring_manager", "")).strip()

        if name:
            paragraphs.append((name, "name"))
        if contact_line:
            paragraphs.append((contact_line, "meta"))
        if date:
            paragraphs.append((date, "meta"))

        paragraphs.append((f"Dear {hiring_manager or 'Hiring Manager'},", "salutation"))

        for key in ("opening_paragraph", "body_paragraph", "closing_paragraph"):
            text = str(context.get(key, "")).strip()
            if text:
                paragraphs.append((text, "body"))

        paragraphs.append(("Sincerely,", "signature"))
        if name:
            paragraphs.append((name, "signature"))
        return paragraphs

    @staticmethod
    def _cover_letter_blank_before(
        paragraphs: list[tuple[str, str]], index: int
    ) -> bool:
        """Return True when a blank line should precede ``paragraphs[index]``.

        Mirrors the ``COVER_LETTER`` template layout: a blank line separates
        the header block from the salutation, every body paragraph, and the
        signature block, while the header (name/contact/date) and the
        signature (``Sincerely,`` + name) each stay contiguous.

        Args:
            paragraphs: The ordered ``(text, kind)`` pairs from
                :meth:`_cover_letter_paragraphs`.
            index: Position of the paragraph to check.

        Returns:
            True when a blank line should be inserted before that paragraph.
        """
        if index == 0:
            return False
        prev_kind = paragraphs[index - 1][1]
        kind = paragraphs[index][1]
        if kind == "body":
            return True
        if kind == "salutation":
            return prev_kind in ("name", "meta")
        if kind == "signature":
            return prev_kind != "signature"
        return False

    @staticmethod
    def _populate_cover_letter_docx(
        doc: DocumentType, context: dict[str, object]
    ) -> None:
        """Populate *doc* with the cover letter content in *context*.

        *context* is the dict produced by :meth:`_build_cover_letter_context`.
        The candidate name is rendered at 14pt bold; every other block uses
        the document's normal style.  A blank paragraph separates the header
        from the salutation, each body paragraph, and the signature block so
        the DOCX letter reads like the plaintext/Markdown versions.
        """
        paragraphs = ResumeRenderer._cover_letter_paragraphs(context)
        for index, (text, kind) in enumerate(paragraphs):
            if ResumeRenderer._cover_letter_blank_before(paragraphs, index):
                doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run(text)
            if kind == "name":
                run.bold = True
                run.font.size = Pt(14)

    @staticmethod
    def _populate_cover_letter_pdf_flowables(
        context: dict[str, object], styles: dict[str, PdfParagraphStyle]
    ) -> list[Flowable]:
        """Build the Platypus flowables for the cover letter in *context*.

        Uses :meth:`_cover_letter_paragraphs` and the shared :meth:`_pdf_styles`
        (name header at 14pt bold, body text otherwise).  A ``Spacer`` the size
        of one body line separates the header from the salutation, each body
        paragraph, and the signature block so the PDF letter reads like the
        plaintext/Markdown versions.  Text is XML-escaped for ReportLab's
        ``Paragraph`` markup parser.
        """
        flowables: list[Flowable] = []
        paragraphs = ResumeRenderer._cover_letter_paragraphs(context)
        for index, (text, kind) in enumerate(paragraphs):
            if ResumeRenderer._cover_letter_blank_before(paragraphs, index):
                flowables.append(Spacer(1, _COVER_LETTER_BLANK_PT))
            style = styles["name"] if kind == "name" else styles["body"]
            flowables.append(Paragraph(escape(text), style))
        return flowables


_COVER_LETTER_BLANK_PT = 13

_SALUTATION_PREFIXES = ("dear", "to whom it may concern", "hello", "hi")
_SIGNATURE_PREFIXES = (
    "sincerely",
    "best regards",
    "kind regards",
    "warm regards",
    "regards",
    "yours truly",
    "yours sincerely",
    "respectfully",
)

_DEFAULT_EXTENSIONS: dict[str, str] = {
    "resume_plaintext": ".txt",
    "resume_markdown": ".md",
    "resume_docx": ".docx",
    "resume_pdf": ".pdf",
    "cover_letter_plaintext": ".txt",
    "cover_letter_markdown": ".md",
    "cover_letter_docx": ".docx",
    "cover_letter_pdf": ".pdf",
    "plaintext": ".txt",
    "markdown": ".md",
    "docx": ".docx",
    "pdf": ".pdf",
}


def _default_extension(document_type: str) -> str:
    """Return the default file extension for *document_type*."""
    return _DEFAULT_EXTENSIONS.get(document_type, ".txt")


def _temp_docx_path() -> Path:
    """Return a fresh temporary ``.docx`` path not tied to the workspace."""
    fd, name = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    return Path(name)


def _temp_pdf_path() -> Path:
    """Return a fresh temporary ``.pdf`` path not tied to the workspace."""
    fd, name = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    return Path(name)


def _bullet_list(items: list[str], styles: dict[str, PdfParagraphStyle]) -> Flowable:
    """Build a bulleted ``ListFlowable`` from *items* using *styles*."""
    flowables = [ListItem(Paragraph(escape(item), styles["bullet"])) for item in items]
    return ListFlowable(cast(Any, flowables), bulletType="bullet")


def _split_letter_body(text: str) -> tuple[list[str], str]:
    """Split *text* into paragraphs plus any trailing contact line.

    The ``COVER_LETTER`` template renders its own salutation and signature,
    so a leading ``Dear ...`` line and a trailing ``Sincerely, ...`` block
    are dropped from the body to avoid duplication.  A signature block may
    carry contact details -- either on its final line (fallback letter) or
    as a separate `` | ``-joined line appended after it (``_apply_contact_info``)
    -- and that contact line is returned so the renderer can place it in the
    letter header instead of discarding it.
    """
    parts = [part.strip() for part in text.split("\n\n") if part.strip()]
    if not parts:
        return parts, ""
    if _is_salutation(parts[0]):
        parts = parts[1:]

    contact_line = ""
    if len(parts) >= 2 and _is_contact_line(parts[-1]) and _is_signature(parts[-2]):
        contact_line = parts.pop()
        parts.pop()
    elif parts and _is_signature(parts[-1]):
        block = parts.pop()
        contact_line = _contact_line_from_signature(block)
    return parts, contact_line


def _is_salutation(paragraph: str) -> bool:
    """Return True when *paragraph* is a standalone greeting line."""
    first_line = paragraph.splitlines()[0].strip().rstrip(",").lower()
    return len(first_line) < 80 and first_line.startswith(_SALUTATION_PREFIXES)


def _is_signature(paragraph: str) -> bool:
    """Return True when *paragraph* is a closing block.

    The first line must be a short sign-off phrase (``Sincerely,``,
    ``Best regards,`` ...).  The block may continue with the candidate's
    name and contact details on later lines, so the length guard applies
    to the first line rather than the whole paragraph.
    """
    first_line = paragraph.splitlines()[0].strip().rstrip(",").lower()
    return len(first_line) < 80 and first_line.startswith(_SIGNATURE_PREFIXES)


def _is_contact_line(paragraph: str) -> bool:
    """Return True when *paragraph* is a standalone contact line."""
    if len(paragraph) > 300 or "\n" in paragraph:
        return False
    return _looks_like_contact(paragraph) and not paragraph.rstrip().endswith(".")


def _contact_line_from_signature(block: str) -> str:
    """Extract a trailing contact line embedded in a signature *block*."""
    lines = [line.strip() for line in block.splitlines() if line.strip()]
    if len(lines) < 3:
        return ""
    contact: list[str] = []
    for line in reversed(lines[1:]):
        if _looks_like_contact(line):
            contact.append(line)
        else:
            break
    return " | ".join(reversed(contact))


def _looks_like_contact(line: str) -> bool:
    """Return True when *line* looks like a contact value or joined line."""
    if " | " in line or "@" in line:
        return True
    if line.startswith(("http://", "https://")):
        return True
    if len(line) >= 40:
        return False
    digits = re.sub(r"\D", "", line)
    return bool(digits) and len(digits) >= 7
